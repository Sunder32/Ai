import os
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
import requests
import json
from typing import List, Optional
import shutil
import subprocess

# Импорт RAG и Fine-tuning модулей
from rag_engine import get_rag_engine
from finetune_prepare import prepare_chat_format, create_ollama_training_modelfile, get_dataset_info
from learning_engine import get_learning_engine

app = FastAPI(title="DeepSeek Mini-Site API")


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

OLLAMA_API_URL = "http://localhost:11434/api/generate"
# Allow overriding the model via env var so you can point to a finetuned model
MODEL_NAME = os.getenv("PROJECT_MODEL_NAME", "deepseek-project-model")
UPLOAD_DIR = "uploads"
DATASET_FILE = "dataset.jsonl"

os.makedirs(UPLOAD_DIR, exist_ok=True)

class ChatRequest(BaseModel):
    prompt: str
    context: Optional[List[int]] = []
    use_learning: Optional[bool] = True  # Использовать контекст обучения

class ChatResponse(BaseModel):
    response: str
    context: List[int]

class CorrectionRequest(BaseModel):
    prompt: str
    original_response: str
    corrected_response: str
    feedback: Optional[str] = ""

class LikeRequest(BaseModel):
    prompt: str
    response: str

@app.post("/api/chat")
async def chat(request: ChatRequest):
    import traceback
    
    try:
        print(f"\n[CHAT] Получен запрос: {request.prompt[:50]}...")
        
        # Получаем контекст обучения если включено
        learning_context = ""
        if request.use_learning:
            learning = get_learning_engine()
            learning_context = learning.get_learning_context(max_examples=5)
        
        # Формируем промпт с контекстом обучения
        enhanced_prompt = request.prompt
        if learning_context:
            enhanced_prompt = f"""{learning_context}

---
Теперь ответь на вопрос пользователя, учитывая исправления выше:

{request.prompt}"""
        
        payload = {
            "model": MODEL_NAME,
            "prompt": enhanced_prompt,
            "stream": False,
            "context": request.context if request.context else []
        }
        
        print(f"[CHAT] Отправка запроса в Ollama с моделью: {MODEL_NAME}")
        
        try:
            response = requests.post(OLLAMA_API_URL, json=payload, timeout=60)
            print(f"[CHAT] Ollama response status: {response.status_code}")
            
            if response.status_code != 200:
                print(f"[CHAT] Ollama error response: {response.text}")
                raise HTTPException(
                    status_code=response.status_code, 
                    detail=f"Ollama returned error: {response.text}"
                )
            
            response.raise_for_status()
            data = response.json()
            
            print(f"[CHAT] Получен ответ от модели: {data.get('response', '')[:50]}...")
            
        except requests.exceptions.Timeout:
            print("[CHAT] ❌ Timeout waiting for Ollama response")
            raise HTTPException(status_code=504, detail="Ollama request timed out")
        except requests.exceptions.ConnectionError:
            print("[CHAT] ❌ Cannot connect to Ollama service")
            raise HTTPException(status_code=503, detail="Ollama service is not reachable. Is it running?")
        
        log_entry = {
            "prompt": request.prompt,
            "response": data.get("response", ""),
            "model": MODEL_NAME
        }
        
        try:
            with open(DATASET_FILE, "a", encoding="utf-8") as f:
                f.write(json.dumps(log_entry, ensure_ascii=False) + "\n")
        except Exception as log_error:
            print(f"[CHAT] Warning: Failed to log conversation: {log_error}")

        return {
            "response": data.get("response", ""),
            "context": data.get("context", [])
        }
        
    except HTTPException:
        raise
    except Exception as e:
        error_details = traceback.format_exc()
        print(f"[CHAT] ❌ Unexpected error:\n{error_details}")
        raise HTTPException(status_code=500, detail=f"Chat failed: {str(e)}")

@app.post("/api/train/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        file_path = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        return {"message": f"File '{file.filename}' uploaded successfully."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/models")
async def list_models():
    try:
        response = requests.get("http://localhost:11434/api/tags")
        response.raise_for_status()
        return response.json()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/train/files")
async def get_uploaded_files():
    try:
        files = []
        if os.path.exists(UPLOAD_DIR):
            for filename in os.listdir(UPLOAD_DIR):
                file_path = os.path.join(UPLOAD_DIR, filename)
                if os.path.isfile(file_path):
                    files.append({
                        "name": filename,
                        "size": os.path.getsize(file_path)
                    })
        return {"files": files, "count": len(files)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/train/start")
async def start_training():
    import traceback
    
    try:
        print("\n" + "="*80)
        print("🚀  НАЧАЛО ОБУЧЕНИЯ НЕЙРОСЕТИ  🚀".center(80))
        print("="*80)
        print(f"[TRAINING] Время старта: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print("-"*80 + "\n")
        
        if not os.path.exists(UPLOAD_DIR) or not os.listdir(UPLOAD_DIR):
            print("[TRAINING]  No files found in upload directory")
            raise HTTPException(status_code=400, detail="No files uploaded for training")
        
        all_text = []
        files_in_dir = os.listdir(UPLOAD_DIR)
        print(f"[TRAINING] Found {len(files_in_dir)} files in upload directory")
        
        for filename in files_in_dir:
            file_path = os.path.join(UPLOAD_DIR, filename)
            if not os.path.isfile(file_path):
                continue
                
            file_ext = os.path.splitext(filename)[1].lower()
            print(f"[TRAINING] Processing file: {filename} (type: {file_ext})")
            
            try:
                if file_ext == '.txt':
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        all_text.append(content)
                        print(f"[TRAINING] Successfully read TXT file: {len(content)} characters")
                
                elif file_ext == '.pdf':
                    try:
                        from pypdf import PdfReader
                        reader = PdfReader(file_path)
                        text = ""
                        for i, page in enumerate(reader.pages):
                            page_text = page.extract_text()
                            text += page_text
                            print(f"[TRAINING] PDF page {i+1}: {len(page_text)} characters")
                        all_text.append(text)
                        print(f"[TRAINING] Successfully read PDF file: {len(text)} characters total")
                    except Exception as pdf_error:
                        print(f"[ERROR] Error processing PDF {filename}: {str(pdf_error)}")
                        traceback.print_exc()
                
                elif file_ext in ['.docx', '.doc']:
                    try:
                        import docx
                        doc = docx.Document(file_path)
                        text = "\n".join([para.text for para in doc.paragraphs])
                        all_text.append(text)
                        print(f"[TRAINING] Successfully read DOCX file: {len(text)} characters")
                    except Exception as docx_error:
                        print(f"[ERROR] Error processing DOCX {filename}: {str(docx_error)}")
                        traceback.print_exc()
                
                elif file_ext == '.json':
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                        
                        # Если это список объектов (датасет)
                        if isinstance(data, list):
                            text_parts = []
                            for item in data:
                                if isinstance(item, dict):
                                    # Извлекаем текстовые поля
                                    for key in ['text', 'content', 'input', 'output', 'question', 'answer', 'instruction']:
                                        if key in item and item[key]:
                                            text_parts.append(str(item[key]))
                                else:
                                    text_parts.append(str(item))
                            text = "\n".join(text_parts)
                        else:
                            text = json.dumps(data, ensure_ascii=False, indent=2)
                        
                        if text.strip():
                            all_text.append(text)
                            print(f"[TRAINING] Successfully read JSON file: {len(text)} characters")
                    except Exception as json_error:
                        print(f"[ERROR] Error processing JSON {filename}: {str(json_error)}")
                        traceback.print_exc()
                
                elif file_ext == '.jsonl':
                    try:
                        text_parts = []
                        with open(file_path, 'r', encoding='utf-8') as f:
                            for line in f:
                                if line.strip():
                                    item = json.loads(line)
                                    if isinstance(item, dict):
                                        for key in ['text', 'content', 'input', 'output', 'prompt', 'response']:
                                            if key in item and item[key]:
                                                text_parts.append(str(item[key]))
                        text = "\n".join(text_parts)
                        if text.strip():
                            all_text.append(text)
                            print(f"[TRAINING] Successfully read JSONL file: {len(text)} characters")
                    except Exception as jsonl_error:
                        print(f"[ERROR] Error processing JSONL {filename}: {str(jsonl_error)}")
                        traceback.print_exc()
                        
            except Exception as e:
                print(f"[ERROR] Error processing {filename}: {str(e)}")
                traceback.print_exc()
                continue
        
        if not all_text:
            print("[ERROR] Could not extract text from any files")
            raise HTTPException(status_code=400, detail="Could not extract text from any files")
        
        combined_text = "\n\n".join(all_text)
        print(f"[TRAINING] Combined text length: {len(combined_text)} characters")
        
        training_file = "training_data.txt"
        with open(training_file, 'w', encoding='utf-8') as f:
            f.write(combined_text)
        print(f"[TRAINING] Saved training data to {training_file}")
        
        context_snippet = combined_text[:6000] if len(combined_text) > 6000 else combined_text
        
        # Escape special characters and format context for SYSTEM prompt
        # Replace newlines with spaces and remove problematic characters
        clean_context = context_snippet.replace('\n', ' ').replace('\r', ' ').replace('"', "'")
        # Remove multiple spaces
        import re
        clean_context = re.sub(r'\s+', ' ', clean_context).strip()
        
        modelfile_content = f'''FROM deepseek-r1:8b
PARAMETER temperature 0.7
PARAMETER top_p 0.9

SYSTEM """You are DeepSeek, an AI assistant with specialized knowledge from the provided documents. Use this knowledge to provide accurate and helpful responses.

Here is your knowledge base context:

{clean_context}

Based on this context and your general knowledge, respond helpfully and accurately to user questions."""'''

        modelfile_path = os.path.abspath("Modelfile")
        with open(modelfile_path, "w", encoding="utf-8") as f:
            f.write(modelfile_content)
        print(f"[TRAINING] Updated Modelfile at: {modelfile_path}")
        
        # Verify Modelfile exists
        if not os.path.exists(modelfile_path):
            raise HTTPException(status_code=500, detail="Failed to create Modelfile")
        
        # Recreate model with Ollama
        print("[TRAINING] Attempting to delete old model...")
        try:
            delete_response = requests.delete(
                "http://localhost:11434/api/delete", 
                json={"name": MODEL_NAME},
                timeout=10
            )
            print(f"[TRAINING] Delete response: {delete_response.status_code}")
        except Exception as del_error:
            print(f"[TRAINING] Delete failed (this is OK if model doesn't exist): {str(del_error)}")
        
        print("[TRAINING] Creating new model...")
        try:
            # Use subprocess to create model via ollama CLI (more reliable)
            print(f"[TRAINING] Using Modelfile at: {modelfile_path}")
            
            # Run ollama create command
            result = subprocess.run(
                ["ollama", "create", MODEL_NAME, "-f", modelfile_path],
                capture_output=True,
                text=True,
                timeout=600  # 10 minutes timeout
            )
            
            print(f"[TRAINING] Ollama create stdout: {result.stdout}")
            if result.stderr:
                print(f"[TRAINING] Ollama create stderr: {result.stderr}")
            
            if result.returncode != 0:
                error_msg = result.stderr or result.stdout or "Unknown error"
                print(f"[ERROR] Failed to create model: {error_msg}")
                raise HTTPException(
                    status_code=500, 
                    detail=f"Failed to create model: {error_msg}"
                )
            
            print(f"[TRAINING] Model '{MODEL_NAME}' created successfully!")
            
        except subprocess.TimeoutExpired:
            print("[ERROR] Timeout while creating model")
            raise HTTPException(status_code=500, detail="Model creation timed out")
        except FileNotFoundError:
            print("[ERROR] Ollama command not found. Make sure Ollama is installed and in PATH")
            raise HTTPException(status_code=503, detail="Ollama is not installed or not in PATH")
        except Exception as create_error:
            print(f"[ERROR] Error creating model: {str(create_error)}")
            raise HTTPException(status_code=500, detail=f"Failed to create model: {str(create_error)}")
        

        print("\n" + "-"*80)
        print(" ОБУЧЕНИЕ ЗАВЕРШЕНО УСПЕШНО!  ✅".center(80))
        print("="*80)
        print(f"[TRAINING] Время завершения: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print(f"[TRAINING] Обработано файлов: {len(all_text)}")
        print(f"[TRAINING] Всего символов: {len(combined_text)}")
        print(f"[TRAINING] Модель: {MODEL_NAME}")
        print("="*80 + "\n")
        
        return {
            "message": "Training completed successfully",
            "files_processed": len(all_text),
            "model_name": MODEL_NAME,
            "total_chars": len(combined_text)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        error_details = traceback.format_exc()
        print(f"[ERROR] Training failed with exception:\n{error_details}")
        raise HTTPException(status_code=500, detail=f"Training failed: {str(e)}")


# ============== RAG API ==============

class RAGSearchRequest(BaseModel):
    query: str
    n_results: int = 5

@app.post("/api/rag/index")
async def rag_index_documents():
    """Индексировать все документы в RAG (векторную БД)"""
    import traceback
    
    try:
        print("\n[RAG] Начало индексации документов...")
        
        if not os.path.exists(UPLOAD_DIR) or not os.listdir(UPLOAD_DIR):
            raise HTTPException(status_code=400, detail="Нет файлов для индексации")
        
        rag = get_rag_engine()
        rag.clear()  # Очищаем старые данные
        
        all_texts = []
        all_metadatas = []
        
        for filename in os.listdir(UPLOAD_DIR):
            file_path = os.path.join(UPLOAD_DIR, filename)
            if not os.path.isfile(file_path):
                continue
            
            file_ext = os.path.splitext(filename)[1].lower()
            text = ""
            
            try:
                if file_ext == '.txt':
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                
                elif file_ext == '.pdf':
                    from pypdf import PdfReader
                    reader = PdfReader(file_path)
                    text = "\n".join(page.extract_text() for page in reader.pages)
                
                elif file_ext in ['.docx', '.doc']:
                    import docx
                    doc = docx.Document(file_path)
                    text = "\n".join(para.text for para in doc.paragraphs)
                
                elif file_ext == '.json':
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    if isinstance(data, list):
                        text = "\n".join(str(item) for item in data)
                    else:
                        text = json.dumps(data, ensure_ascii=False)
                
                if text.strip():
                    all_texts.append(text)
                    all_metadatas.append({"filename": filename})
                    print(f"[RAG] Обработан файл: {filename} ({len(text)} символов)")
                    
            except Exception as e:
                print(f"[RAG] Ошибка обработки {filename}: {e}")
                continue
        
        if not all_texts:
            raise HTTPException(status_code=400, detail="Не удалось извлечь текст из файлов")
        
        # Индексируем
        chunks_count = rag.add_documents(all_texts, all_metadatas)
        
        print(f"[RAG] ✅ Индексация завершена: {chunks_count} чанков")
        
        return {
            "message": "Индексация завершена",
            "files_processed": len(all_texts),
            "chunks_created": chunks_count
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[RAG] ❌ Ошибка: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/rag/search")
async def rag_search(request: RAGSearchRequest):
    """Поиск по документам через RAG"""
    try:
        rag = get_rag_engine()
        results = rag.search(request.query, request.n_results)
        
        return {
            "query": request.query,
            "results": results
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/rag/stats")
async def rag_stats():
    """Статистика RAG базы"""
    try:
        rag = get_rag_engine()
        return rag.get_stats()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/chat/rag")
async def chat_with_rag(request: ChatRequest):
    """Чат с использованием RAG (поиск контекста перед ответом)"""
    import traceback
    
    try:
        print(f"\n[RAG-CHAT] Получен запрос: {request.prompt[:50]}...")
        
        # Получаем релевантный контекст из RAG
        rag = get_rag_engine()
        context = rag.get_context_for_query(request.prompt, max_tokens=2000)
        
        # Формируем промпт с контекстом
        if context:
            enhanced_prompt = f"""Используй следующий контекст для ответа на вопрос:

КОНТЕКСТ:
{context}

ВОПРОС: {request.prompt}

Отвечай на основе контекста. Если информации нет в контексте, скажи об этом."""
        else:
            enhanced_prompt = request.prompt
        
        payload = {
            "model": MODEL_NAME,
            "prompt": enhanced_prompt,
            "stream": False,
            "context": request.context if request.context else []
        }
        
        response = requests.post(OLLAMA_API_URL, json=payload, timeout=60)
        response.raise_for_status()
        data = response.json()
        
        return {
            "response": data.get("response", ""),
            "context": data.get("context", []),
            "rag_context_used": bool(context)
        }
        
    except Exception as e:
        print(f"[RAG-CHAT] ❌ Ошибка: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))


# ============== FINE-TUNING API ==============

@app.post("/api/finetune/prepare")
async def prepare_finetune_data():
    """Подготовить данные для fine-tuning из JSON и TXT файлов"""
    try:
        print("\n[FINETUNE] Подготовка данных для fine-tuning...")
        
        prepared_files = []
        
        for filename in os.listdir(UPLOAD_DIR):
            file_ext = os.path.splitext(filename)[1].lower()
            
            # Поддерживаем JSON и TXT файлы
            if file_ext not in ['.json', '.txt']:
                continue
                
            file_path = os.path.join(UPLOAD_DIR, filename)
            
            try:
                # Получаем информацию о датасете
                info = get_dataset_info(file_path)
                print(f"[FINETUNE] Датасет {filename}: {info}")
                
                if "error" in info:
                    print(f"[FINETUNE] Пропуск {filename}: {info['error']}")
                    continue
                
                # Подготавливаем данные в зависимости от типа файла
                if file_ext == '.json':
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    
                    if not isinstance(data, list):
                        print(f"[FINETUNE] Пропуск {filename}: не является списком")
                        continue
                    
                    output_path = prepare_chat_format(data, file_path)
                    
                elif file_ext == '.txt':
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    
                    output_path = prepare_chat_format(text, file_path)
                
                prepared_files.append({
                    "source": filename,
                    "output": os.path.basename(output_path),
                    "info": info
                })
                
                print(f"[FINETUNE] ✅ Подготовлен: {output_path}")
                
            except Exception as e:
                print(f"[FINETUNE] Ошибка обработки {filename}: {e}")
                import traceback
                traceback.print_exc()
                continue
        
        if not prepared_files:
            raise HTTPException(status_code=400, detail="Не найдено JSON или TXT файлов с данными для fine-tuning")
        
        return {
            "message": "Данные подготовлены для fine-tuning",
            "files": prepared_files
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/finetune/create-model")
async def create_finetuned_model():
    """Создать модель с few-shot примерами из датасета"""
    try:
        print("\n[FINETUNE] Создание модели с примерами...")
        
        # Ищем JSON или TXT датасет
        dataset_path = None
        for filename in os.listdir(UPLOAD_DIR):
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext in ['.json', '.txt']:
                dataset_path = os.path.join(UPLOAD_DIR, filename)
                break
        
        if not dataset_path:
            raise HTTPException(status_code=400, detail="Не найден JSON или TXT датасет")
        
        print(f"[FINETUNE] Используем датасет: {dataset_path}")
        
        # Создаём Modelfile с примерами
        modelfile_path = os.path.abspath("Modelfile.finetune")
        result_path = create_ollama_training_modelfile(
            base_model="deepseek-r1:8b",
            dataset_path=dataset_path,
            output_path=modelfile_path
        )
        
        print(f"[FINETUNE] Создан Modelfile: {result_path}")
        
        # Создаём модель в Ollama
        model_name = "deepseek-finetuned"
        
        result = subprocess.run(
            ["ollama", "create", model_name, "-f", modelfile_path],
            capture_output=True,
            text=True,
            timeout=300
        )
        
        if result.returncode != 0:
            error_msg = result.stderr or result.stdout or "Unknown error"
            print(f"[FINETUNE] ❌ Ошибка Ollama: {error_msg}")
            raise HTTPException(status_code=500, detail=f"Ошибка создания модели: {error_msg}")
        
        print(f"[FINETUNE] ✅ Модель {model_name} создана!")
        
        return {
            "message": f"Модель {model_name} успешно создана",
            "model_name": model_name,
            "modelfile": modelfile_path
        }
        
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"[FINETUNE] ❌ Ошибка: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/finetune/info")
async def get_finetune_info():
    """Информация о датасетах для fine-tuning (JSON и TXT)"""
    try:
        datasets = []
        
        if not os.path.exists(UPLOAD_DIR):
            return {"datasets": [], "message": "Папка uploads не существует"}
        
        for filename in os.listdir(UPLOAD_DIR):
            file_ext = os.path.splitext(filename)[1].lower()
            
            # Поддерживаем JSON, TXT и JSONL
            if file_ext not in ['.json', '.txt', '.jsonl']:
                continue
                
            file_path = os.path.join(UPLOAD_DIR, filename)
            
            try:
                info = get_dataset_info(file_path)
                info["filename"] = filename
                datasets.append(info)
            except Exception as e:
                print(f"[FINETUNE] Ошибка чтения {filename}: {e}")
                datasets.append({
                    "filename": filename,
                    "error": str(e)
                })
        
        return {"datasets": datasets, "count": len(datasets)}
        
    except Exception as e:
        import traceback
        print(f"[FINETUNE] ❌ Ошибка get_finetune_info: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))


# ============== LEARNING API (Самообучение) ==============

@app.post("/api/learning/correct")
async def submit_correction(request: CorrectionRequest):
    """
    Отправить исправление ответа ИИ
    ИИ запомнит ошибку и не будет её повторять
    """
    try:
        learning = get_learning_engine()
        result = learning.add_correction(
            prompt=request.prompt,
            original_response=request.original_response,
            corrected_response=request.corrected_response,
            feedback=request.feedback or ""
        )
        
        print(f"[LEARNING] ✅ Добавлено исправление: {request.prompt[:50]}...")
        
        return {
            "message": "Исправление сохранено! ИИ учтёт это в будущих ответах.",
            "correction_id": result["id"]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/learning/like")
async def like_response(request: LikeRequest):
    """
    Лайкнуть хороший ответ ИИ
    ИИ будет стараться отвечать похожим образом
    """
    try:
        learning = get_learning_engine()
        result = learning.add_good_response(
            prompt=request.prompt,
            response=request.response
        )
        
        print(f"[LEARNING] 👍 Ответ помечен как хороший")
        
        return {
            "message": "Спасибо за отзыв! ИИ будет учиться на хороших примерах.",
            "id": result["id"]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/learning/stats")
async def get_learning_stats():
    """Статистика обучения"""
    try:
        learning = get_learning_engine()
        return learning.get_stats()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/learning/export")
async def export_learning_data():
    """Экспортировать данные обучения для fine-tuning"""
    try:
        learning = get_learning_engine()
        examples = []
        
        # Добавляем исправления
        for corr in learning.corrections:
            examples.append({
                "messages": [
                    {"role": "user", "content": corr['prompt']},
                    {"role": "assistant", "content": corr['corrected_response']}
                ]
            })
        
        # Добавляем хорошие ответы
        for good in learning.good_responses:
            examples.append({
                "messages": [
                    {"role": "user", "content": good['prompt']},
                    {"role": "assistant", "content": good['response']}
                ]
            })
        
        return {
            "examples": examples,
            "count": len(examples)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============== FILE RESPONSES ==============

@app.post("/api/chat/file")
async def chat_with_file_response(request: ChatRequest):
    """
    Чат с возможностью ответа файлом (Markdown, код и т.п.)
    """
    import traceback
    import re
    
    try:
        # Добавляем инструкцию для файловых ответов
        enhanced_prompt = f"""Ты можешь отвечать в формате Markdown. 
Используй:
- Заголовки (# ## ###)
- Списки (- или 1.)
- Код (```язык ... ```)
- Таблицы
- Выделение (**жирный**, *курсив*)

Если пользователь просит создать файл, код или документ — оформи это в соответствующем формате.

Вопрос пользователя: {request.prompt}"""

        # Получаем контекст обучения
        learning = get_learning_engine()
        learning_context = learning.get_learning_context(max_examples=3)
        
        if learning_context:
            enhanced_prompt = f"{learning_context}\n\n---\n\n{enhanced_prompt}"
        
        payload = {
            "model": MODEL_NAME,
            "prompt": enhanced_prompt,
            "stream": False,
            "context": request.context if request.context else []
        }
        
        response = requests.post(OLLAMA_API_URL, json=payload, timeout=120)
        response.raise_for_status()
        data = response.json()
        
        response_text = data.get("response", "")
        
        # Определяем тип контента
        content_type = "text"
        file_extension = None
        
        # Проверяем, есть ли блоки кода
        code_blocks = re.findall(r'```(\w+)?\n(.*?)```', response_text, re.DOTALL)
        
        if code_blocks:
            lang = code_blocks[0][0] or "txt"
            content_type = "code"
            file_extension = {
                "python": "py", "javascript": "js", "typescript": "ts",
                "html": "html", "css": "css", "json": "json",
                "markdown": "md", "md": "md", "sql": "sql",
                "bash": "sh", "shell": "sh", "yaml": "yaml", "yml": "yaml"
            }.get(lang.lower(), lang.lower())
        
        # Проверяем, похоже ли на Markdown документ
        if re.search(r'^#{1,3}\s', response_text, re.MULTILINE):
            content_type = "markdown"
            file_extension = "md"
        
        return {
            "response": response_text,
            "context": data.get("context", []),
            "content_type": content_type,
            "file_extension": file_extension,
            "can_download": content_type in ["code", "markdown"]
        }
        
    except Exception as e:
        print(f"[CHAT-FILE] ❌ Ошибка: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))


app.mount("/", StaticFiles(directory="client", html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", "5050"))
    uvicorn.run(app, host="0.0.0.0", port=port)
